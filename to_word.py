from docx import Document


def to_word(jobs, filename='job_listings.docx'):
    """
    Generates a Word document with job listings.

    :param jobs: List of job dictionaries. Each dictionary should have 'Title', 'Link', and 'Description'.
    :param filename: Name of the Word document file to save (default is 'job_listings.docx').
    """
    # Create a new Document
    doc = Document()

    # Add a title for the document
    doc.add_heading('Job Listings', level=1)

    # Iterate over the jobs list and add each job to the document
    for job in jobs:
        # Add job title as a heading
        doc.add_heading(job['Title'], level=2)

        # Add job link
        doc.add_paragraph(f'Link: {job["Link"]}')

        # Add job description
        doc.add_paragraph(f'Description: {job["Description"]}')

        # Add a horizontal line between jobs
        doc.add_paragraph('-----------------------------------')

    # Save the document
    doc.save(filename)
    print(f"Word document '{filename}' created successfully.")


# # Example usage
# jobs = [
#     {'Title': 'Java Developer Level 2',
#      'Link': 'https://www.kforce.com/find-work/search-jobs/#/detail/MTY5Nn5FUUd-MjExNzgzN1Qxfjk5/',
#      'Description': 'Kforce has a client in Cincinnati, OH that is seeking a Level 2 Java Developer...'},
#     {'Title': 'Python Developer',
#      'Link': 'https://example.com/job/python-developer',
#      'Description': 'This position requires strong Python skills...'},
#     # Add more job dictionaries as needed
# ]
#
# generate_job_listings_doc(jobs)
